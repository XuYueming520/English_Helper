# This py file is for writing files of English helper by Xu Yueming.
"""@By 徐跃鸣@"""
from inspect import currentframe
from itertools import chain
from os import listdir
from os.path import isfile
from pickle import load, dump

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Pt, RGBColor
# 导包

def easy_save(input, path=None):
    """简易保存
        input -> 任何类型 : 要保存的数据
        path -> str或None : 路径"""
    if path is None:  # 没有指定保存路径
        flag = False  # 是否找到变量
        fr = currentframe().f_back
        for var_name, value in chain(fr.f_locals.items(), fr.f_globals.items()):  # 遍历所有对象
            if value is input:  # 获取input实参变量名var_name
                path = var_name  # 路径设为变量名
                flag = True  # 找到变量
                break  # 退出循环
        if not flag:  # 还是没有找到
            path = 'Unknown'  # 设置成默认值"Unknown"
    dump(input, open(f'{path}.EasySave', 'wb'))  # 保存,默认加上后缀名".EasySave"

def easy_load(path=None, mode=1):
    """简易读取
        path -> str或list或None : 路径
        mode -> int : 是否修改全局变量"""
    if type(path) == str:  # 字符类型
        path = [path]  # 统一成列表
    elif path is None:  # 没有指定
        path = []  # 初始化清空
        for f in listdir('.\\'):  # 遍历程序跟目下所有文件
            if f.endswith('.EasySave'):  # 如果后缀名是默认保存的
                path.append(f)  # 添加文件
    dictionary = {}  # 初始化
    for file in path:  # 遍历需要读取的文件
        info = load(open(file, 'rb'))  # 读取文件
        var_name = '.'.join(file.split('.')[:-1])  # 变量名称(文件名去除后缀名)
        if mode == 0:  # 模式0是直接修改全局变量(!有时候这会对程序造成严重的Bug甚至安全漏洞!)
            globals()[var_name] = info  # 修改
        dictionary[var_name] = info  # 修改
    return dictionary  # 返回

class Docx_writer(object):
    """docx文件操作"""
    def __init__(self):
        """初始化"""
        self.docx_file = Document()  # 创建新的文件
        
        self.dict = ['⁰', '¹', '²', '³', '⁴', '⁵', '⁶', '⁷', '⁸', '⁹']  # 行号，列号
        self.docx_width = 432  # word文档可用宽度(A4)
    
    def set_size(self, original, size):
        """设置字体大小
            original -> docx.text.run.Run : 原始文字
            size -> int : 大小(磅值)"""
        original.font.size = Pt(size)  # 设置大小
    
    def set_color(self, original, color):
        """设置颜色
            original -> docx.text.run.Run : 原始文字
            color -> RGBColor或list或tuple : 颜色"""
        if type(color) in [list, tuple]:  # 不是RGBColor(即为list或tuple)
            color = RGBColor(*color)  # 转换成RGBColor
        original.font.color.rgb = color  # 设置颜色
    
    def set_font(self, original, font):
        """设置字体
            original -> docx.text.run.Run : 原始文字
            color -> str : 字体"""
        original.font.name = font
        original._element.rPr.rFonts.set(qn('w:eastAsia'), font)  # 同上一行设置字体
    
    def write_cross_puzzle(self, puzzle, sentences, name='徐跃鸣', where_from='Module 5201314',
                           title='Crossword puzzle', NO=False, hide=None, NOcolor='ff8080', words_color='000000',
                           Chinese_font='宋体', English_font='arial'):
        """写入crossword puzzle
            puzzle -> list : 要保存的puzzle
            sentences -> list : 要保存的句子
            name -> str : 作者姓名
            where_from -> str : 单词来源
            title -> str : 大标题
            NO -> AllType : puzzle中代表空的值
            hide -> AllType : puzzle中单词不显示的值
            NOcolor -> str : puzzle中空的单元格的颜色
            words_color -> str : 文字颜色
            Chinese_font -> str : 中文字体
            English_font -> str : 英文字体"""
        self.docx_file = Document()  # 创建新的文件
        
        lv = len(words_color)
        words_color = tuple(int(words_color[i:i + lv // 3], 16) for i in range(0, lv, lv // 3))  # 将十六进制的颜色值转化成RGB颜色值
        
        title_paragraph = self.docx_file.add_paragraph('')  # 新增标题段落
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 设置标题
        t = title_paragraph.add_run(title)  # 标题
        t.bold = True  # 粗体
        self.set_size(t, 20)  # 设置大小
        self.set_color(t, words_color)  # 设置颜色
        self.set_font(t, English_font)  # 设置字体
        
        info_paragraph = self.docx_file.add_paragraph('')  # 新增信息段落
        wfm = info_paragraph.add_run('Words from Module ')  # 添加提示信息单词来源
        self.set_size(wfm, 12)  # 设置大小
        self.set_color(wfm, words_color)  # 设置颜色
        self.set_font(wfm, English_font)  # 设置字体
        wf = info_paragraph.add_run(where_from)  # 单词来源
        wf.underline = True  # 下划线
        self.set_size(wf, 12)  # 设置大小
        self.set_color(wf, words_color)  # 设置颜色
        self.set_font(wf, English_font)  # 设置字体
        cb = info_paragraph.add_run('                      Created by ')  # 添加提示制作人
        self.set_size(cb, 12)  # 设置大小
        self.set_color(cb, words_color)  # 设置颜色
        self.set_font(cb, English_font)  # 设置字体
        n = info_paragraph.add_run(name)  # 制作人
        n.underline = True  # 下划线
        self.set_size(n, 12)  # 设置大小
        self.set_color(n, words_color)  # 设置颜色
        self.set_font(n, Chinese_font)  # 设置字体

        R = len(puzzle[0])  # 行数
        C = len(puzzle)  # 列数
        a = self.docx_width // C - 9  # 每个单元格的边长
        
        table = self.docx_file.add_table(rows=len(puzzle[0]), cols=len(puzzle), style='Table Grid')  # 表格
        self.set_font(table.style, English_font)  # 设置字体
        self.set_color(table.style, words_color)  # 设置颜色
        self.set_size(table.style, a)  # 设置大小
        
        cnt = 1  # DeBug
        
        for r in range(R):  # 遍历列
            for c in range(C):  # 遍历行
                run = table.cell(r, c).paragraphs[0]
                run.alignment = WD_ALIGN_PARAGRAPH.LEFT
                width = a  # 宽度
                height = a  # 高度
                if r == 0:  # 第一行
                    if c <= 9:  # 个位数
                        width += 15  # 数字的空间
                    else:  # 有十位
                        width += 20  # 数字的空间
                        height += 10
                    x = run.add_run(''.join([self.dict[int(x)] for x in str(c + 1)]) + '\n')  # 需要添加数字表示现在是第几列
                    self.set_size(x, a)
                elif c == 0:  # 第一列,但不是第一行,避免(0,0)位置有两个数字
                    if c <= 9:  # 个位数
                        width += 15  # 数字的空间
                    else:  # 有十位
                        width += 30  # 数字的空间
                    x = run.add_run(''.join([self.dict[int(x)] for x in str(r + 1)]))  # 需要添加数字表示现在是第几行
                    self.set_size(x, a)
                if puzzle[c][r] != hide:  # 不是隐藏的字母
                    if puzzle[c][r] != NO:  # 是字母
                        if len(puzzle[c][r]) > 1:  # 有大小写区分
                            width += 40  # 宽度增加
                        run.add_run(puzzle[c][r])  # 放入字母
                    else:  # 是空白部分
                        table.rows[r].cells[c]._tc.get_or_add_tcPr().append(parse_xml((r'<w:shd {} w:fill="' + NOcolor + r'"/>').format(nsdecls('w'))))  # 设置空白部分的颜色，区别于隐藏单词
                table.cell(r, c).width = width
                table.cell(r, c).height = height  # 设置单元格宽度,高度
                print("\r%.2f%%" % (cnt * 100.0 / (C * R)), end='')  # DeBug
                cnt += 1  # DeBug
                
        CROSS = self.docx_file.add_paragraph('')  # 提示信息句子 横行
        DOWN = self.docx_file.add_paragraph('')  # 提示信息句子 纵行
        
        c = CROSS.add_run('Cross:\n')  # 横向提示句子
        self.set_font(c, English_font)  # 设置字体
        self.set_color(c, words_color)  # 设置颜色
        self.set_size(c, 12)  # 设置大小
        for (x, y) in sentences[0].items():  # 遍历横向的句子
            _ = CROSS.add_run(f'    {x}:{y}\n')
            self.set_font(_, English_font)  # 设置字体
            self.set_color(_, words_color)  # 设置颜色
            self.set_size(_, 12)  # 设置大小
        
        c = DOWN.add_run('Down:\n')  # 竖向提示句子
        self.set_font(c, English_font)  # 设置字体
        self.set_color(c, words_color)  # 设置颜色
        self.set_size(c, 12)  # 设置大小
        for (x, y) in sentences[1].items():  # 遍历竖向的句子
            _ = DOWN.add_run(f'    {x}:{y}\n')
            self.set_font(_, English_font)  # 设置字体
            self.set_color(_, words_color)  # 设置颜色
            self.set_size(_, 12)  # 设置大小
    
    def save(self, file_path=r'.\Crossword puzzle.docx', overwrite=False):
        """保存
            file_path -> str : 文件保存的路径
            overwrite -> bool : 是否覆盖"""
        if not overwrite and isfile(file_path):  # 如果文件已经存在
            print(f'{file_path} 已经存在文件,保存失败!')  # 输出提示
            if input("是否覆盖(Y,N)>>>") == 'N':  # 询问是否覆盖
                return -1  # 返回错误值
        self.docx_file.save(file_path)  # 保存
        return 0  # 返回正常值

if __name__ == '__main__':  # DeBug
    docx_test = Docx_writer()
    docx_test.write_cross_puzzle([['a'] * 25] * 25, [{'0': '我喜欢'}, {'1': '你'}])
    docx_test.save(r'D:\Xu\Desktop\a.docx', overwrite=True)